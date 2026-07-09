"""
output_tool.py - Word 文件輸出工具（修正版）
修正：所有輸出到 Word 前，先用 force_to_text() 統一轉字串
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime

from utils import force_to_text, force_list_of_str


def create_word_document(all_results: list, output_path: str) -> str:
    """將所有結構化內容整合成一個完整的 Word 文件"""
    doc = Document()

    # 標題
    title = doc.add_heading('保險員證照考試教材整理', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para = doc.add_paragraph(f'生成時間：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 統計摘要
    doc.add_heading('📊 處理統計', level=1)
    total_files = len(all_results)
    total_pages = sum(len(r.get('structured_pages', [])) for r in all_results)
    cross_pages = sum(
        1 for r in all_results
        for p in r.get('structured_pages', [])
        if p.get('is_cross_page')
    )
    flowcharts = sum(
        1 for r in all_results
        for p in r.get('structured_pages', [])
        if p.get('has_flowchart')
    )
    tables_count = sum(
        1 for r in all_results
        for p in r.get('structured_pages', [])
        if p.get('tables')
    )

    stats = doc.add_paragraph()
    stats.add_run(f'檔案數量：{total_files} 個\n')
    stats.add_run(f'總頁數：{total_pages} 頁\n')
    stats.add_run(f'跨頁合併：{cross_pages} 處\n')
    stats.add_run(f'流程圖頁面：{flowcharts} 頁\n')
    stats.add_run(f'含表格頁面：{tables_count} 頁\n')
    doc.add_page_break()

    # 依檔案整理內容
    for file_result in all_results:
        file_name = force_to_text(file_result.get('file_name', '未知檔案'))
        pages = file_result.get('structured_pages', [])

        doc.add_heading(f'📄 {file_name}', level=1)

        for page_data in pages:
            page_num = page_data.get('page_num', 0)
            main_topic = force_to_text(page_data.get('main_topic', '未知主題'))
            is_cross = bool(page_data.get('is_cross_page', False))
            has_flowchart = bool(page_data.get('has_flowchart', False))

            # 頁碼標題
            page_title = f'第 {page_num} 頁 — {main_topic}'
            if is_cross:
                page_title += ' 【跨頁合併】'
            if has_flowchart:
                page_title += ' 【含流程圖】'
            doc.add_heading(page_title, level=2)

            # 流程圖步驟
            flowchart_steps = force_list_of_str(page_data.get('flowchart_steps', []))
            flowchart_steps = [s for s in flowchart_steps if s.strip()]
            if flowchart_steps:
                doc.add_heading('🔄 流程圖步驟', level=3)
                flow_para = doc.add_paragraph()
                flow_run = flow_para.add_run(' → '.join(flowchart_steps))
                flow_run.bold = True
                flow_run.font.size = Pt(11)
                flow_run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)

            # 重點整理
            key_points = force_list_of_str(page_data.get('key_points', []))
            key_points = [p for p in key_points if p.strip()]
            if key_points:
                doc.add_heading('📌 重點整理', level=3)
                for point in key_points:
                    para = doc.add_paragraph(style='List Bullet')
                    para.add_run(force_to_text(point)).font.size = Pt(11)

            # 專業術語
            definitions = page_data.get('definitions', []) or []
            valid_defs = [d for d in definitions if isinstance(d, dict) and d]
            if valid_defs:
                doc.add_heading('📖 專業術語', level=3)
                for defn in valid_defs:
                    term = force_to_text(defn.get('term', ''))
                    definition = force_to_text(defn.get('definition', ''))
                    if term or definition:
                        para = doc.add_paragraph()
                        run_term = para.add_run(f'{term}：')
                        run_term.bold = True
                        run_term.font.size = Pt(11)
                        para.add_run(definition).font.size = Pt(11)

            # 表格內容
            tables_data = page_data.get('tables', []) or []
            valid_tables = [t for t in tables_data if isinstance(t, dict) and t]
            if valid_tables:
                doc.add_heading('📋 表格內容', level=3)
                for tbl in valid_tables:
                    title_text = force_to_text(tbl.get('title', '表格'))
                    content_text = force_to_text(tbl.get('content', ''))
                    para = doc.add_paragraph()
                    para.add_run(f'{title_text}：').bold = True
                    if content_text:
                        doc.add_paragraph(content_text)

            # 考試重點（紅色）
            exam_hints = force_list_of_str(page_data.get('exam_hints', []))
            exam_hints = [h for h in exam_hints if h.strip()]
            if exam_hints:
                doc.add_heading('⭐ 考試重點', level=3)
                for hint in exam_hints:
                    para = doc.add_paragraph(style='List Bullet')
                    run = para.add_run(force_to_text(hint))
                    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                    run.font.size = Pt(11)

            # 完整結構化內容
            structured_content = force_to_text(page_data.get('structured_content', ''))
            if structured_content.strip():
                doc.add_heading('📝 完整內容', level=3)
                doc.add_paragraph(structured_content)

            doc.add_paragraph('─' * 50)

        doc.add_page_break()

    doc.save(output_path)
    return output_path
